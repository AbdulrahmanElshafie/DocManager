"""
Tests for DOCX to PDF conversion functionality
"""
import os
import tempfile
from django.test import TestCase
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from rest_framework.test import APIClient
from rest_framework import status

from manager.models import Document
from manager.utils.convert_docx_to_pdf import convert_docx_to_pdf, DocxToPdfConverter

User = get_user_model()


class DocxToPdfConversionTest(TestCase):
    """Test cases for DOCX to PDF conversion utility"""
    
    def setUp(self):
        # Create a test user
        self.user = User.objects.create_user(
            username='testuser',
            email='test@example.com',
            password='testpassword123'
        )
        self.client = APIClient()
        self.client.force_authenticate(user=self.user)
    
    def create_test_docx_file(self) -> SimpleUploadedFile:
        """Create a simple test DOCX file"""
        from docx import Document as DocxDocument
        
        # Create a simple DOCX file
        doc = DocxDocument()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test DOCX document for PDF conversion.')
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            doc.save(temp_file.name)
            
            # Read the file content
            with open(temp_file.name, 'rb') as f:
                content = f.read()
            
            # Clean up
            os.unlink(temp_file.name)
            
            return SimpleUploadedFile(
                name='test_document.docx',
                content=content,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    
    def test_convert_docx_bytes_to_pdf(self):
        """Test converting DOCX bytes to PDF"""
        # Create test DOCX file
        docx_file = self.create_test_docx_file()
        
        try:
            # Convert to PDF
            pdf_bytes = convert_docx_to_pdf(docx_file.read())
            
            # Verify PDF content
            self.assertIsInstance(pdf_bytes, bytes)
            self.assertGreater(len(pdf_bytes), 0)
            
            # Check PDF magic bytes
            self.assertTrue(pdf_bytes.startswith(b'%PDF-'))
            
        except RuntimeError as e:
            # If LibreOffice is not available, the fallback should still work
            # This test should not be skipped as we have a working fallback
            self.fail(f"Conversion failed even with fallback: {e}")
    
    def test_convert_docx_file_to_pdf(self):
        """Test converting DOCX file path to PDF"""
        # Create test DOCX file
        docx_file = self.create_test_docx_file()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(docx_file.read())
            temp_file.flush()
            
            try:
                # Convert to PDF
                pdf_bytes = convert_docx_to_pdf(temp_file.name)
                
                # Verify PDF content
                self.assertIsInstance(pdf_bytes, bytes)
                self.assertGreater(len(pdf_bytes), 0)
                self.assertTrue(pdf_bytes.startswith(b'%PDF-'))
                
            except RuntimeError as e:
                # Should work with fallback
                self.fail(f"Conversion failed even with fallback: {e}")
            finally:
                # Clean up
                os.unlink(temp_file.name)
    
    def test_invalid_file_extension(self):
        """Test that non-DOCX files raise ValueError"""
        # Create a text file with .docx extension
        with tempfile.NamedTemporaryFile(delete=False, suffix='.txt') as temp_file:
            temp_file.write(b'This is not a DOCX file')
            temp_file.flush()
            
            try:
                with self.assertRaises(ValueError):
                    convert_docx_to_pdf(temp_file.name)
            finally:
                os.unlink(temp_file.name)
    
    def test_file_not_found(self):
        """Test that missing files raise ValueError"""
        with self.assertRaises(ValueError):
            convert_docx_to_pdf('/nonexistent/file.docx')


class DocumentPdfConversionAPITest(TestCase):
    """Test cases for the Document PDF conversion API endpoint"""
    
    def setUp(self):
        # Create a test user
        self.user = User.objects.create_user(
            username='testuser',
            email='test@example.com',
            password='testpassword123'
        )
        self.client = APIClient()
        self.client.force_authenticate(user=self.user)
    
    def create_test_docx_file(self) -> SimpleUploadedFile:
        """Create a simple test DOCX file"""
        from docx import Document as DocxDocument
        
        # Create a simple DOCX file
        doc = DocxDocument()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test DOCX document for PDF conversion.')
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            doc.save(temp_file.name)
            
            # Read the file content
            with open(temp_file.name, 'rb') as f:
                content = f.read()
            
            # Clean up
            os.unlink(temp_file.name)
            
            return SimpleUploadedFile(
                name='test_document.docx',
                content=content,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    
    def test_convert_docx_document_to_pdf(self):
        """Test the API endpoint for converting DOCX to PDF"""
        # Create a test DOCX document
        docx_file = self.create_test_docx_file()
        
        # Create a Document object
        document = Document.objects.create(
            name='Test Document',
            file=docx_file,
            owner=self.user
        )
        
        # Make request to convert to PDF
        url = f'/api/manager/document/{document.id}/convert/pdf/'
        response = self.client.get(url)
        
        # Check response - should work with our fallback
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertIn('inline; filename="Test Document.pdf"', response['Content-Disposition'])
        
        # Check PDF content
        pdf_content = response.content
        self.assertIsInstance(pdf_content, bytes)
        self.assertGreater(len(pdf_content), 0)
        self.assertTrue(pdf_content.startswith(b'%PDF-'))
    
    def test_convert_non_docx_document_returns_error(self):
        """Test that non-DOCX documents return appropriate error"""
        # Create a test PDF document (not DOCX)
        pdf_content = b'%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n'
        pdf_file = SimpleUploadedFile(
            name='test_document.pdf',
            content=pdf_content,
            content_type='application/pdf'
        )
        
        # Create a Document object
        document = Document.objects.create(
            name='Test PDF Document',
            file=pdf_file,
            owner=self.user
        )
        
        # Make request to convert to PDF
        url = f'/api/manager/document/{document.id}/convert/pdf/'
        response = self.client.get(url)
        
        # Check response
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        response_data = response.json()
        self.assertIn('not a DOCX document', response_data['error'])
    
    def test_convert_nonexistent_document(self):
        """Test converting a document that doesn't exist"""
        # Use a non-existent document ID
        url = '/api/manager/document/00000000-0000-0000-0000-000000000000/convert/pdf/'
        response = self.client.get(url)
        
        # Should return 404
        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)
    
    def test_convert_document_without_file(self):
        """Test converting a document without an associated file"""
        # Create a Document object without a file
        document = Document.objects.create(
            name='Test Document',
            owner=self.user
        )
        
        # Make request to convert to PDF
        url = f'/api/manager/document/{document.id}/convert/pdf/'
        response = self.client.get(url)
        
        # Check response
        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)
        response_data = response.json()
        self.assertEqual(response_data['error'], 'File not found')