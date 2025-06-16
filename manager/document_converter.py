import os
import tempfile
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown2
import html2text
from bs4 import BeautifulSoup
import re
from typing import Tuple, Union


class DocumentConverter:
    """Utility class for converting between HTML/Markdown and DOCX formats"""
    
    @staticmethod
    def docx_to_html(docx_file_path: str) -> str:
        """Convert a DOCX file to HTML format"""
        try:
            doc = Document(docx_file_path)
            html_content = ['<html><body>']
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Basic formatting - detect headings, bold, etc.
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading ', '')
                        try:
                            level = int(level)
                            html_content.append(f'<h{level}>{text}</h{level}>')
                        except ValueError:
                            html_content.append(f'<h1>{text}</h1>')
                    else:
                        # Handle basic text formatting
                        formatted_text = DocumentConverter._apply_formatting(text, paragraph)
                        html_content.append(f'<p>{formatted_text}</p>')
                else:
                    html_content.append('<p></p>')  # Empty paragraph
            
            # Handle tables
            for table in doc.tables:
                html_content.append('<table border="1">')
                for row in table.rows:
                    html_content.append('<tr>')
                    for cell in row.cells:
                        html_content.append(f'<td>{cell.text}</td>')
                    html_content.append('</tr>')
                html_content.append('</table>')
            
            html_content.append('</body></html>')
            return '\n'.join(html_content)
            
        except Exception as e:
            # If conversion fails, return a basic HTML structure
            return f'<html><body><p>Error converting document: {str(e)}</p></body></html>'
    
    @staticmethod
    def _apply_formatting(text: str, paragraph) -> str:
        """Apply basic formatting to text based on paragraph runs"""
        try:
            formatted_text = ""
            for run in paragraph.runs:
                run_text = run.text
                if run.bold:
                    run_text = f'<strong>{run_text}</strong>'
                if run.italic:
                    run_text = f'<em>{run_text}</em>'
                if run.underline:
                    run_text = f'<u>{run_text}</u>'
                formatted_text += run_text
            return formatted_text if formatted_text else text
        except:
            return text
    
    @staticmethod
    def html_to_docx(html_content: str, output_path: str) -> bool:
        """Convert HTML content to DOCX format"""
        try:
            doc = Document()
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style tags
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Process HTML elements
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'br', 'table']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(element.name[1])
                    paragraph = doc.add_heading(element.get_text().strip(), level=level)
                elif element.name == 'p' or element.name == 'div':
                    text = element.get_text().strip()
                    if text:
                        paragraph = doc.add_paragraph()
                        DocumentConverter._add_formatted_text(paragraph, element)
                elif element.name == 'br':
                    doc.add_paragraph()
                elif element.name == 'table':
                    DocumentConverter._add_table(doc, element)
            
            # If no content was added, add a blank paragraph
            if len(doc.paragraphs) == 0:
                doc.add_paragraph("Document content")
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error converting HTML to DOCX: {e}")
            # Create a simple document with error message
            try:
                doc = Document()
                doc.add_paragraph(f"Error converting document: {str(e)}")
                doc.save(output_path)
                return True
            except:
                return False
    
    @staticmethod
    def _add_formatted_text(paragraph, element):
        """Add formatted text to a paragraph from HTML element"""
        try:
            for item in element.contents:
                if hasattr(item, 'name'):  # It's a tag
                    run = paragraph.add_run(item.get_text())
                    if item.name == 'strong' or item.name == 'b':
                        run.bold = True
                    elif item.name == 'em' or item.name == 'i':
                        run.italic = True
                    elif item.name == 'u':
                        run.underline = True
                else:  # It's text
                    paragraph.add_run(str(item))
        except:
            # Fallback to plain text
            paragraph.add_run(element.get_text())
    
    @staticmethod
    def _add_table(doc, table_element):
        """Add a table to the document from HTML table element"""
        try:
            rows = table_element.find_all('tr')
            if not rows:
                return
            
            # Count columns from first row
            cols = len(rows[0].find_all(['td', 'th']))
            if cols == 0:
                return
            
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Light Grid Accent 1'
            
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells[:cols]):  # Ensure we don't exceed column count
                    table.cell(i, j).text = cell.get_text().strip()
        except Exception as e:
            print(f"Error adding table: {e}")
    
    @staticmethod
    def markdown_to_html(markdown_content: str) -> str:
        """Convert Markdown content to HTML"""
        try:
            # Use markdown2 for conversion with extra features
            html = markdown2.markdown(markdown_content, extras=[
                'fenced-code-blocks',
                'tables',
                'strike',
                'task_list',
                'wiki-tables',
                'code-friendly'
            ])
            return f'<html><body>{html}</body></html>'
        except Exception as e:
            print(f"Error converting Markdown to HTML: {e}")
            return f'<html><body><p>{markdown_content}</p></body></html>'
    
    @staticmethod
    def html_to_markdown(html_content: str) -> str:
        """Convert HTML content to Markdown"""
        try:
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = False
            h.body_width = 0  # Don't wrap lines
            return h.handle(html_content)
        except Exception as e:
            print(f"Error converting HTML to Markdown: {e}")
            # Fallback to text extraction
            soup = BeautifulSoup(html_content, 'html.parser')
            return soup.get_text()
    
    @staticmethod
    def get_document_format(file_path: str) -> str:
        """Determine the format of a document file"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.docx':
            return 'docx'
        elif ext == '.html':
            return 'html'
        elif ext in ['.md', '.markdown']:
            return 'markdown'
        else:
            return 'unknown'
    
    @staticmethod
    def convert_for_editor(docx_file_path: str) -> Tuple[str, str]:
        """Convert DOCX to HTML for editor and return both HTML and Markdown"""
        html_content = DocumentConverter.docx_to_html(docx_file_path)
        markdown_content = DocumentConverter.html_to_markdown(html_content)
        return html_content, markdown_content
    
    @staticmethod
    def convert_from_editor(content, content_type, file_path):
        """Convert content from editor format back to original document format"""
        try:
            if file_path.endswith('.docx'):
                return DocumentConverter._save_to_docx(content, content_type, file_path)
            elif file_path.endswith('.html'):
                return DocumentConverter._save_to_html(content, content_type, file_path)
            elif file_path.endswith('.md'):
                return DocumentConverter._save_to_markdown(content, content_type, file_path)
            else:
                # For other formats, save as plain text
                with open(file_path, 'w', encoding='utf-8') as f:
                    if content_type == 'html':
                        # Convert HTML to plain text
                        h = html2text.HTML2Text()
                        h.ignore_links = True
                        text_content = h.handle(content)
                        f.write(text_content)
                    else:
                        f.write(content)
                return True
        except Exception as e:
            print(f"Error converting document: {e}")
            return False
    
    @staticmethod
    def _save_to_docx(content, content_type, file_path):
        """Save content to DOCX format"""
        try:
            doc = Document()
            
            if content_type == 'html':
                # Parse HTML and convert to DOCX
                soup = BeautifulSoup(content, 'html.parser')
                
                # Remove html and body tags if present
                if soup.html:
                    soup = soup.html
                if soup.body:
                    soup = soup.body
                
                for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li']):
                    text = element.get_text().strip()
                    if not text:
                        continue
                    
                    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        # Add heading
                        level = int(element.name[1])
                        heading = doc.add_heading(text, level=level)
                    elif element.name == 'p':
                        # Add paragraph
                        doc.add_paragraph(text)
                    elif element.name in ['ul', 'ol']:
                        # Handle lists
                        for li in element.find_all('li'):
                            doc.add_paragraph(li.get_text().strip(), style='List Bullet')
            
            elif content_type == 'markdown':
                # Convert markdown to HTML first, then to DOCX
                html_content = markdown2.markdown(content)
                return DocumentConverter._save_to_docx(html_content, 'html', file_path)
            
            else:
                # Plain text
                lines = content.split('\n')
                for line in lines:
                    if line.strip():
                        doc.add_paragraph(line.strip())
                    else:
                        doc.add_paragraph('')
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            doc.save(file_path)
            return True
            
        except Exception as e:
            print(f"Error saving DOCX: {e}")
            return False
    
    @staticmethod
    def _save_to_html(content, content_type, file_path):
        """Save content to HTML format"""
        try:
            if content_type == 'markdown':
                html_content = markdown2.markdown(content)
            elif content_type == 'html':
                html_content = content
            else:
                # Plain text to HTML
                html_content = f"<html><body>{content.replace(chr(10), '<br>')}</body></html>"
            
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            return True
            
        except Exception as e:
            print(f"Error saving HTML: {e}")
            return False
    
    @staticmethod
    def _save_to_markdown(content, content_type, file_path):
        """Save content to Markdown format"""
        try:
            if content_type == 'html':
                # Convert HTML to markdown
                h = html2text.HTML2Text()
                h.ignore_links = False
                markdown_content = h.handle(content)
            elif content_type == 'markdown':
                markdown_content = content
            else:
                # Plain text (just save as is)
                markdown_content = content
            
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            return True
            
        except Exception as e:
            print(f"Error saving Markdown: {e}")
            return False 